import os
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Cm
from docx.text.paragraph import Paragraph
from PIL import Image
from openpyxl.styles import Font

# =====================================================================
# 【模組 A】Word 自動化文件生成系統 (SOP/規格書)
# =====================================================================

def load_data_from_excel(excel_path):
    """從 Excel 讀取資料，並自動清理 Key 欄位的雜訊"""
    print(f"  [Word] 正在讀取 Excel 資料來源: {excel_path} (分頁: 文字與圖片) ...")
    
    # 讀取文字資料分頁
    df_text = pd.read_excel(excel_path, sheet_name="文字資料填報")
    df_text['變數名稱 (Key)'] = df_text['變數名稱 (Key)'].astype(str).str.strip().str.replace('{', '').str.replace('}', '')
    text_data = dict(zip(df_text['變數名稱 (Key)'], df_text['目前填入值 (工程師填寫)']))
    
    # 讀取圖片資料分頁
    df_img = pd.read_excel(excel_path, sheet_name="圖片路徑填報")
    df_img['圖片變數名稱 (Key)'] = df_img['圖片變數名稱 (Key)'].astype(str).str.strip().str.replace('{', '').str.replace('}', '')
    image_data = dict(zip(df_img['圖片變數名稱 (Key)'], df_img['圖片檔案路徑 (工程師填寫)']))
    
    return text_data, image_data

def process_paragraph_safely(paragraph, text_data, image_data):
    """【隔離級別安全引擎 V3】處理文字與圖片替換，含NaN留白防呆"""
    if not paragraph.text:
        return

    has_text_placeholder = any(f"{{{{{k}}}}}" in paragraph.text for k in text_data.keys() if not pd.isna(k))
    has_img_placeholder = any(f"{{{{IMG_{k}}}}}" in paragraph.text for k in image_data.keys() if not pd.isna(k))
    
    if not (has_text_placeholder or has_img_placeholder):
        return

    text_runs = []
    combined_text = ""
    
    for run in paragraph.runs:
        if run._r.xpath('.//*[local-name()="drawing"]'):
            continue
        text_runs.append(run)
        combined_text += run.text

    if not text_runs:
        return

    for r in text_runs:
        r.text = ""
        
    first_run = text_runs[0]

    # 1. 執行文字精準替換
    for key, value in text_data.items():
        if pd.isna(key): continue
        actual_value = "" if pd.isna(value) else str(value)
        placeholder = f"{{{{{key}}}}}"
        
        if placeholder in combined_text:
            combined_text = combined_text.replace(placeholder, actual_value)
            print(f"    - [成功] 替換標籤: {placeholder} -> '{actual_value}'")

    # 2. 執行圖片標籤抹除與路徑記錄
    img_to_insert = []
    for key, img_path in image_data.items():
        if pd.isna(key): continue
        placeholder = f"{{{{IMG_{key}}}}}"
        
        if placeholder in combined_text:
            combined_text = combined_text.replace(placeholder, "")
            if not pd.isna(img_path) and isinstance(img_path, str):
                if os.path.exists(img_path):
                    img_to_insert.append((key, img_path))
                else:
                    print(f"    - [路徑警告] 找不到圖片檔案: {img_path}")

    first_run.text = combined_text

    # 3. 執行圖片插入 (動態尺寸判斷 + 特定圖片專屬尺寸)
    DEFAULT_MAX_W = 6.5
    DEFAULT_MAX_H = 5.0
    CUSTOM_IMG_RULES = {
        "prodToTest": {"max_w": 7.3, "max_h": 5.3},
    }

    for key, img_path in img_to_insert:
        try:
            with Image.open(img_path) as img:
                orig_width_px, orig_height_px = img.size
            
            aspect_ratio = orig_width_px / orig_height_px

            if key in CUSTOM_IMG_RULES:
                limit_w = CUSTOM_IMG_RULES[key]["max_w"]
                limit_h = CUSTOM_IMG_RULES[key]["max_h"]
            else:
                limit_w = DEFAULT_MAX_W
                limit_h = DEFAULT_MAX_H

            theory_height_cm = limit_w / aspect_ratio
            new_run = paragraph.add_run()
            
            if theory_height_cm <= limit_h:
                new_run.add_picture(img_path, width=Cm(limit_w))
            else:
                new_run.add_picture(img_path, height=Cm(limit_h))

        except Exception as img_err:
            print(f"    - [圖片插入失敗] 處理圖片 {img_path} 時發生錯誤: {img_err}")

def scan_blocks(container, text_data, image_data, doc):
    """【AI 遞迴級別全境掃描引擎】掃描段落、表格、圖形文字"""
    if hasattr(container, 'paragraphs'):
        for p in container.paragraphs:
            process_paragraph_safely(p, text_data, image_data)
            
    if hasattr(container, 'tables'):
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    scan_blocks(cell, text_data, image_data, doc)

    element = None
    if hasattr(container, '_element'):
        element = container._element
    elif hasattr(container, 'element'):
        element = container.element

    if element is not None:
        for txbx in element.xpath('.//*[local-name()="txbxContent"]'):
            for p_elem in txbx.xpath('.//*[local-name()="p"]'):
                p_obj = Paragraph(p_elem, doc)
                process_paragraph_safely(p_obj, text_data, image_data)

def generate_report(template_path, output_path, text_data, image_data):
    if not os.path.exists(template_path):
        print(f"\n[錯誤] 找不到 Word 範本: {template_path}")
        return

    print("  [Word] 開始掃描與生成 Word 文件...")
    doc = Document(template_path)
    
    scan_blocks(doc, text_data, image_data, doc)
    
    for section in doc.sections:
        headers = [section.header, section.first_page_header, section.even_page_header]
        footers = [section.footer, section.first_page_footer, section.even_page_footer]
        
        for h in headers:
            if h is not None: scan_blocks(h, text_data, image_data, doc)
        for f in footers:
            if f is not None: scan_blocks(f, text_data, image_data, doc)
                
    doc.save(output_path)
    print(f"  [Word 大成功] 規格報告已安全生成: {output_path}")


# =====================================================================
# 【模組 B】MES 自動化路由表生成系統 (Excel)
# =====================================================================

def find_cell_location(df, target_value):
    """動態尋找 DataFrame 中特定字串的(行, 列)索引"""
    for r_idx in range(df.shape[0]):
        for c_idx in range(df.shape[1]):
            if str(df.iat[r_idx, c_idx]).strip() == str(target_value).strip():
                return r_idx, c_idx
    return None, None

def generate_mes_files(checklist_path, output_path):
    print(f"  [MES] 正在讀取 Excel 資料來源: {checklist_path} (分頁: fixed/var table) ...")
    
    df_fixed = pd.read_excel(checklist_path, sheet_name='MES_fixed table', header=None)
    df_var = pd.read_excel(checklist_path, sheet_name='MES_var table', header=None)
    
    SHEET_COLUMNS = {
        'Machine': ['SBU', 'MODEL_NO', 'MODEL_NAME', 'SAP_ROUTE', 'GROUP_CODE', 'GROUP_NAME', 'DEVICE_TYPE_CODE', 'DEVICE_TYPE_NAME', 'ACOMD_QTY', 'DEL_FLAG'],
        'Soft': ['SBU', 'MODEL_NO', 'MODEL_NAME', 'SAP_ROUTE', 'GROUP_CODE', 'GROUP_NAME', 'SOFT_NO', 'SOFT_NAME', 'DEL_FLAG'],
        'SOP': ['SBU', 'MODEL_NO', 'MODEL_NAME', 'SAP_ROUTE', 'GROUP_CODE', 'GROUP_NAME', 'SOP_NO', 'SOP_NAME', 'DEL_FLAG'],
        'Tool': ['SBU', 'MODEL_NO', 'MODEL_NAME', 'SAP_ROUTE', 'GROUP_CODE', 'GROUP_NAME', 'TOOL_MODEL', 'TOOL_MODEL_NAME', 'USE_QTY', 'DEL_FLAG'],
        'Route': ['SBU', 'MODEL_NO', 'MODEL_NAME', 'SAP_ROUTE', 'LEFT', 'GROUP_CODE', 'GROUP_NAME', 'FACTORY_CODE', 'FACTORY_NAME', 'BASE_QTY', 'SPLIT_QTY', 'SETUP_TIME1', 'SETUP_TIME2', 'LABOR_TIME1', 'LABOR_TIME2', 'MACHINE_TIME', 'STD_TEXT_NO']
    }
    
    output_sheets = {sheet: pd.DataFrame(columns=cols) for sheet, cols in SHEET_COLUMNS.items()}

    # 第一階段：提取共用資訊
    r_sbu, c_sbu = find_cell_location(df_fixed, "SBU")
    sbu_val = df_fixed.iat[r_sbu + 1, c_sbu] if r_sbu is not None else None

    r_model_no, c_model_no = find_cell_location(df_var, "MODEL_NO")
    model_no_val = df_var.iat[r_model_no + 1, c_model_no] if r_model_no is not None else None

    r_model_name, c_model_name = find_cell_location(df_var, "MODEL_NAME")
    model_name_val = df_var.iat[r_model_name + 1, c_model_name] if r_model_name is not None else None

    r_sap, c_sap = find_cell_location(df_fixed, "SAP_ROUTE")
    sap_routes = []
    if r_sap is not None:
        for i in range(1, 4):
            if r_sap + i < df_fixed.shape[0]:
                val = df_fixed.iat[r_sap + i, c_sap]
                if pd.notna(val): sap_routes.append(val)
                     
    r_fac_code, c_fac_code = find_cell_location(df_fixed, "FACTORY_CODE")
    factory_codes = [df_fixed.iat[r_fac_code+i, c_fac_code] for i in range(1,4)] if r_fac_code is not None else [None]*3
    
    r_fac_name, c_fac_name = find_cell_location(df_fixed, "FACTORY_NAME")
    factory_names = [df_fixed.iat[r_fac_name+i, c_fac_name] for i in range(1,4)] if r_fac_name is not None else [None]*3

    group_map = {}
    r_gcode, c_gcode = find_cell_location(df_fixed, "GROUP_CODE")
    r_gname, c_gname = find_cell_location(df_fixed, "GROUP_NAME")
    r_stno, c_stno = find_cell_location(df_fixed, "STD_TEXT_NO")
    
    if r_gcode is not None and r_gname is not None:
        r = r_gcode + 1
        while r < df_fixed.shape[0]:
            gn = df_fixed.iat[r, c_gname]
            if pd.isna(gn) or str(gn).strip() == '': break
            gc = df_fixed.iat[r, c_gcode]
            st = df_fixed.iat[r, c_stno] if c_stno is not None else None
            group_map[str(gn).strip()] = {'code': gc, 'std_text': st}
            r += 1

    # 第二階段：各分頁資料組裝
    print("  [MES] 正在組裝各分頁資料...")
    # Soft 分頁
    r_sg, c_sg = find_cell_location(df_var, "GROUP_NAME")
    r_sno, c_sno = find_cell_location(df_var, "SOFT_NO")
    soft_data_rows = []
    if r_sg is not None and r_sno is not None:
        r = r_sg + 1
        while r < df_var.shape[0]:
            gn = str(df_var.iat[r, c_sg]).strip()
            if gn == "DONE": break
            if pd.isna(df_var.iat[r, c_sg]) or gn in ['', 'nan', 'None']: 
                r += 1
                continue
            
            soft_no = df_var.iat[r, c_sno]
            soft_name = df_var.iat[r, c_sno + 1] 
            if pd.notna(soft_no) and str(soft_no).strip() not in ['', 'nan', 'None']:
                 for route in sap_routes: 
                     g_code = group_map.get(gn, {}).get('code', None)
                     soft_data_rows.append({
                         'SBU': sbu_val, 'MODEL_NO': model_no_val, 'MODEL_NAME': model_name_val,
                         'SAP_ROUTE': route, 'GROUP_CODE': g_code, 'GROUP_NAME': gn,
                         'SOFT_NO': soft_no, 'SOFT_NAME': soft_name, 'DEL_FLAG': None
                     })
            r += 1
    output_sheets['Soft'] = pd.DataFrame(soft_data_rows, columns=SHEET_COLUMNS['Soft'])

    # SOP 分頁
    r_sopno, c_sopno = find_cell_location(df_var, "SOP_NO")
    sop_data_rows = []
    if r_sg is not None and r_sopno is not None:
        r = r_sg + 1
        while r < df_var.shape[0]:
            gn = str(df_var.iat[r, c_sg]).strip()
            if gn == "DONE": break
            if pd.isna(df_var.iat[r, c_sg]) or gn in ['', 'nan', 'None']: 
                r += 1
                continue
            
            sop_no = df_var.iat[r, c_sopno]
            sop_name = df_var.iat[r, c_sopno + 1] 
            if pd.notna(sop_no) and str(sop_no).strip() not in ['', 'nan', 'None']:
                 for route in sap_routes:
                     g_code = group_map.get(gn, {}).get('code', None)
                     sop_data_rows.append({
                         'SBU': sbu_val, 'MODEL_NO': model_no_val, 'MODEL_NAME': model_name_val,
                         'SAP_ROUTE': route, 'GROUP_CODE': g_code, 'GROUP_NAME': gn,
                         'SOP_NO': sop_no, 'SOP_NAME': sop_name, 'DEL_FLAG': None
                     })
            r += 1
    output_sheets['SOP'] = pd.DataFrame(sop_data_rows, columns=SHEET_COLUMNS['SOP'])

    # Tool 分頁
    r_tool, c_tool = find_cell_location(df_var, "TOOL_MODEL")
    tool_data_rows = []
    if r_sg is not None and r_tool is not None:
        r = r_sg + 1
        while r < df_var.shape[0]:
            gn = str(df_var.iat[r, c_sg]).strip()
            if gn == "DONE": break
            if pd.isna(df_var.iat[r, c_sg]) or gn in ['', 'nan', 'None']: 
                r += 1
                continue
            
            tool_model = df_var.iat[r, c_tool]
            tool_name = df_var.iat[r, c_tool + 1]
            use_qty = df_var.iat[r, c_tool + 2]
            if pd.notna(tool_model) and str(tool_model).strip() not in ['', 'nan', 'None']:
                 for route in sap_routes:
                     g_code = group_map.get(gn, {}).get('code', None)
                     tool_data_rows.append({
                         'SBU': sbu_val, 'MODEL_NO': model_no_val, 'MODEL_NAME': model_name_val,
                         'SAP_ROUTE': route, 'GROUP_CODE': g_code, 'GROUP_NAME': gn,
                         'TOOL_MODEL': tool_model, 'TOOL_MODEL_NAME': tool_name,
                         'USE_QTY': use_qty, 'DEL_FLAG': None
                     })
            r += 1
    output_sheets['Tool'] = pd.DataFrame(tool_data_rows, columns=SHEET_COLUMNS['Tool'])

    # Route 分頁
    unique_groups = []
    if r_sg is not None:
        r = r_sg + 1
        while r < df_var.shape[0]:
            gn = str(df_var.iat[r, c_sg]).strip()
            if gn == "DONE": break
            if pd.isna(df_var.iat[r, c_sg]) or gn in ['', 'nan', 'None']: 
                r += 1
                continue
            if gn not in unique_groups:
                 unique_groups.append(gn)
            r += 1

    route_data_rows = []
    for i, route in enumerate(sap_routes):
        fac_code = factory_codes[i] if i < len(factory_codes) else None
        fac_name = factory_names[i] if i < len(factory_names) else None
        
        current_left = 200
        for gn in unique_groups:
            g_code = group_map.get(gn, {}).get('code', None)
            std_text = group_map.get(gn, {}).get('std_text', None) 
            
            route_data_rows.append({
                'SBU': sbu_val, 'MODEL_NO': model_no_val, 'MODEL_NAME': model_name_val,
                'SAP_ROUTE': route, 'LEFT': current_left, 'GROUP_CODE': g_code,
                'GROUP_NAME': gn, 'FACTORY_CODE': fac_code, 'FACTORY_NAME': fac_name,
                'BASE_QTY': 1, 'SPLIT_QTY': 0, 'SETUP_TIME1': 300,
                'SETUP_TIME2': 0, 'LABOR_TIME1': 0, 'LABOR_TIME2': 0,
                'MACHINE_TIME': 0, 'STD_TEXT_NO': std_text
            })
            current_left += 150

    output_sheets['Route'] = pd.DataFrame(route_data_rows, columns=SHEET_COLUMNS['Route'])

    # 第三階段：寫出最終檔案並套用紅字樣式
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        for sheet_name in ['Machine', 'Soft', 'SOP', 'Tool', 'Route']:
            output_sheets[sheet_name].to_excel(writer, sheet_name=sheet_name, index=False)
        
        workbook = writer.book
        if 'Route' in workbook.sheetnames:
            route_sheet = workbook['Route']
            red_font = Font(color="FF0000")
            target_red_columns = ['BASE_QTY', 'SPLIT_QTY', 'SETUP_TIME1', 'SETUP_TIME2', 'LABOR_TIME1', 'LABOR_TIME2', 'MACHINE_TIME']
            red_col_indices = []
            
            for col_idx in range(1, route_sheet.max_column + 1):
                if route_sheet.cell(row=1, column=col_idx).value in target_red_columns:
                    red_col_indices.append(col_idx)
            
            for row_idx in range(2, route_sheet.max_row + 1):
                for col_idx in red_col_indices:
                    route_sheet.cell(row=row_idx, column=col_idx).font = red_font

    print(f"  [MES 大成功] MES 路由表已生成，數值與紅色格式已套用: {output_path}")

# =====================================================================
# 【主程式】統一執行區塊
# =====================================================================

if __name__ == "__main__":
    # --- 共同來源檔案 ---
    EXCEL_CHECKLIST = "Checklist_Template.xlsx"

    # --- 任務一：Word 產出參數 ---
    WORD_TEMPLATE = "制式文件範本.docx"
    FINAL_WORD_OUTPUT = "最終輸出_新產品規格報告.docx"
    
    # --- 任務二：MES Excel 產出參數 ---
    FINAL_EXCEL_OUTPUT = "MES route.xlsx"
    
    print("=" * 60)
    print("【啟動雙效自動化流程】")
    print("=" * 60)

    # 執行任務一：產出 Word
    print("\n▶ 開始執行任務一：產出 Word 規格報告...")
    try:
        t_data, i_data = load_data_from_excel(EXCEL_CHECKLIST)
        generate_report(WORD_TEMPLATE, FINAL_WORD_OUTPUT, t_data, i_data)
    except Exception as e:
        print(f"  [Word 系統錯誤] 執行過程中發生異常: {e}")

    # 執行任務二：產出 MES Excel
    print("\n▶ 開始執行任務二：產出 MES 路由 Excel 檔案...")
    try:
        generate_mes_files(EXCEL_CHECKLIST, FINAL_EXCEL_OUTPUT)
    except Exception as e:
        print(f"  [MES 系統錯誤] 執行過程中發生異常: {e}")
        
    print("\n" + "=" * 60)
    print("【全自動化流程完畢】")
    print("=" * 60)